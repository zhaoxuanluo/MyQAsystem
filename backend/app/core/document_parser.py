"""Multi-format document parser.

Supports: PDF, DOCX, CSV, TXT, MD
Each parser returns a list of DocumentPage objects with text and metadata.
"""

import csv
import io
from dataclasses import dataclass, field
from pathlib import Path

import fitz  # PyMuPDF
import markdown
from docx import Document as DocxDocument

try:
    import chardet
except ImportError:  # pragma: no cover - optional dependency
    chardet = None


@dataclass
class DocumentPage:
    """A single page/section extracted from a document."""
    text: str
    page_number: int | None = None
    section_title: str | None = None
    metadata: dict = field(default_factory=dict)


class DocumentParser:
    """Parse various document formats into structured text pages."""

    PARSERS = {
        "pdf": "_parse_pdf",
        "docx": "_parse_docx",
        "doc": "_parse_docx",
        "csv": "_parse_csv",
        "txt": "_parse_txt",
        "md": "_parse_md",
    }

    def parse(self, file_path: str) -> list[DocumentPage]:
        """Parse a file and return a list of pages."""
        ext = Path(file_path).suffix.lstrip(".").lower()
        parser_method = self.PARSERS.get(ext)
        if not parser_method:
            raise ValueError(f"Unsupported file type: {ext}")
        return getattr(self, parser_method)(file_path)

    def _parse_pdf(self, file_path: str) -> list[DocumentPage]:
        """Extract text from PDF using PyMuPDF."""
        pages = []
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text")
            if text.strip():
                pages.append(DocumentPage(
                    text=text.strip(),
                    page_number=page_num,
                    metadata={"source": Path(file_path).name},
                ))
        doc.close()
        return pages

    def _parse_docx(self, file_path: str) -> list[DocumentPage]:
        """Extract text from DOCX preserving heading structure."""
        doc = DocxDocument(file_path)
        pages = []
        current_section = None
        current_text_parts: list[str] = []
        page_counter = 1

        for para in doc.paragraphs:
            # Detect headings as section boundaries
            if para.style.name.startswith("Heading"):
                # Save previous section
                if current_text_parts:
                    pages.append(DocumentPage(
                        text="\n".join(current_text_parts).strip(),
                        page_number=page_counter,
                        section_title=current_section,
                        metadata={"source": Path(file_path).name},
                    ))
                    page_counter += 1
                    current_text_parts = []
                current_section = para.text.strip()
                current_text_parts.append(para.text)
            elif para.text.strip():
                current_text_parts.append(para.text)

        # Save last section
        if current_text_parts:
            pages.append(DocumentPage(
                text="\n".join(current_text_parts).strip(),
                page_number=page_counter,
                section_title=current_section,
                metadata={"source": Path(file_path).name},
            ))

        # If no headings found, return as single page
        if not pages:
            full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if full_text:
                pages.append(DocumentPage(
                    text=full_text,
                    page_number=1,
                    metadata={"source": Path(file_path).name},
                ))

        return pages

    def _parse_csv(self, file_path: str) -> list[DocumentPage]:
        """Convert CSV rows to readable text paragraphs."""
        pages = []
        content, used_encoding = self._read_text_file(file_path)
        if content is None:
            raise ValueError("Unable to decode CSV file with supported encodings")

        # Parse CSV
        reader = csv.DictReader(io.StringIO(content))
        rows = list(reader)

        if not rows:
            return pages

        # Convert CSV to structured text in batches of 20 rows
        headers = list(rows[0].keys())
        batch_size = 20

        for i in range(0, len(rows), batch_size):
            batch = rows[i:i + batch_size]
            text_parts = []
            for row in batch:
                row_text = "; ".join(f"{k}: {v}" for k, v in row.items() if v)
                text_parts.append(row_text)

            pages.append(DocumentPage(
                text="\n".join(text_parts),
                page_number=(i // batch_size) + 1,
                section_title=f"Data rows {i + 1}-{i + len(batch)}",
                metadata={
                    "source": Path(file_path).name,
                    "headers": headers,
                    "total_rows": len(rows),
                    "encoding": used_encoding,
                },
            ))

        return pages

    def _parse_txt(self, file_path: str) -> list[DocumentPage]:
        """Parse plain text file."""
        text, used_encoding = self._read_text_file(file_path)
        if text is None:
            raise ValueError("Unable to decode text file with supported encodings")

        if not text.strip():
            return []

        # Split by double newlines for natural paragraph grouping
        sections = text.split("\n\n")
        pages = []
        current_text_parts: list[str] = []
        page_counter = 1
        char_count = 0

        for section in sections:
            section = section.strip()
            if not section:
                continue
            current_text_parts.append(section)
            char_count += len(section)

            # Group ~2000 chars per page
            if char_count >= 2000:
                pages.append(DocumentPage(
                    text="\n\n".join(current_text_parts),
                    page_number=page_counter,
                    metadata={
                        "source": Path(file_path).name,
                        "encoding": used_encoding,
                    },
                ))
                current_text_parts = []
                char_count = 0
                page_counter += 1

        if current_text_parts:
            pages.append(DocumentPage(
                text="\n\n".join(current_text_parts),
                page_number=page_counter,
                metadata={
                    "source": Path(file_path).name,
                    "encoding": used_encoding,
                },
            ))

        return pages

    def _parse_md(self, file_path: str) -> list[DocumentPage]:
        """Parse Markdown file preserving heading structure."""
        text, used_encoding = self._read_text_file(file_path)
        if text is None:
            raise ValueError("Unable to decode markdown file with supported encodings")

        if not text.strip():
            return []

        # Split by top-level headings
        lines = text.split("\n")
        pages = []
        current_section = None
        current_text_parts: list[str] = []
        page_counter = 1

        for line in lines:
            if line.startswith("# ") or line.startswith("## "):
                # Save previous section
                if current_text_parts:
                    pages.append(DocumentPage(
                        text="\n".join(current_text_parts).strip(),
                        page_number=page_counter,
                        section_title=current_section,
                        metadata={
                            "source": Path(file_path).name,
                            "encoding": used_encoding,
                        },
                    ))
                    page_counter += 1
                    current_text_parts = []
                current_section = line.lstrip("# ").strip()
            current_text_parts.append(line)

        if current_text_parts:
            pages.append(DocumentPage(
                text="\n".join(current_text_parts).strip(),
                page_number=page_counter,
                section_title=current_section,
                metadata={
                    "source": Path(file_path).name,
                    "encoding": used_encoding,
                },
            ))

        return pages

    def _read_text_file(self, file_path: str) -> tuple[str | None, str | None]:
        """Read a text file with broad encoding fallback."""
        raw = Path(file_path).read_bytes()

        encodings: list[str] = []
        if chardet is not None:
            detected = chardet.detect(raw)
            if detected.get("encoding"):
                encodings.append(detected["encoding"])

        encodings.extend([
            "utf-8-sig",
            "utf-8",
            "utf-16",
            "utf-16-le",
            "utf-16-be",
            "gb18030",
            "gbk",
            "gb2312",
            "big5",
            "cp936",
            "latin-1",
        ])

        seen: set[str] = set()
        for encoding in encodings:
            if not encoding:
                continue
            encoding = encoding.lower()
            if encoding in seen:
                continue
            seen.add(encoding)
            try:
                return raw.decode(encoding), encoding
            except UnicodeDecodeError:
                continue

        # Final fallback: decode with replacement so uploads do not fail outright
        return raw.decode("utf-8", errors="replace"), "utf-8(replace)"


# Singleton
document_parser = DocumentParser()
